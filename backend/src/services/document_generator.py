"""
Single-module document generator for ATS-friendly PDF and DOCX resumes.
Replaces the dependency on external Open Resume service with a self-contained solution.

Features:
- Clean, professional ATS-optimized templates
- Support for both PDF and DOCX formats
- Single-column layout for better ATS parsing
- Standard fonts and proper formatting
- Direct integration with Gemini-tailored resume data
"""

from __future__ import annotations

import io
from datetime import datetime
from typing import List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from loguru import logger
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

try:
    from ..models.resume import Resume
except ImportError:
    from models.resume import Resume


class DocumentGeneratorService:
    """
    Unified service for generating ATS-friendly resumes in PDF and DOCX formats.
    Uses a clean, single-column template optimized for Applicant Tracking Systems.
    """

    # ATS-friendly configuration
    FONT_NAME = "Helvetica"  # Standard font available in ReportLab
    FONT_SIZE_NAME = 18
    FONT_SIZE_SECTION = 14
    FONT_SIZE_BODY = 11
    FONT_SIZE_SMALL = 10

    COLOR_PRIMARY = colors.HexColor("#1a1a1a")  # Near-black for text
    COLOR_SECTION = colors.HexColor("#2563eb")  # Professional blue for headers
    COLOR_LINK = colors.HexColor("#1d4ed8")  # Darker blue for links

    MARGIN = 0.75 * inch
    LINE_SPACING = 1.15

    def __init__(self):
        """Initialize the document generator service"""
        logger.info("Initialized DocumentGeneratorService")

    # ==================== PDF GENERATION ====================

    def generate_pdf(self, resume: Resume) -> bytes:
        """
        Generate an ATS-friendly PDF resume.

        Args:
            resume: Resume model from Gemini tailoring

        Returns:
            PDF file as bytes

        Raises:
            Exception: If PDF generation fails
        """
        try:
            logger.info(
                f"Generating PDF for resume: {resume.name} (candidate: {resume.personalInfo.name})"
            )

            # Create PDF buffer
            buffer = io.BytesIO()

            # Create PDF document
            doc = SimpleDocTemplate(
                buffer,
                pagesize=letter,
                rightMargin=self.MARGIN,
                leftMargin=self.MARGIN,
                topMargin=self.MARGIN,
                bottomMargin=self.MARGIN,
                title=f"{resume.personalInfo.name} - Resume",
                author=resume.personalInfo.name,
            )

            # Build content
            story = []
            styles = self._create_pdf_styles()

            # Header section
            self._add_pdf_header(story, styles, resume)

            # Summary section (if present)
            if resume.personalInfo.summary:
                self._add_pdf_section(story, styles, "PROFESSIONAL SUMMARY")
                summary_para = Paragraph(
                    self._sanitize_text(resume.personalInfo.summary), styles["Body"]
                )
                story.append(summary_para)
                story.append(Spacer(1, 0.2 * inch))

            # Experience section
            if resume.experience:
                self._add_pdf_section(story, styles, "PROFESSIONAL EXPERIENCE")
                for exp in resume.experience:
                    self._add_pdf_experience(story, styles, exp)

            # Education section
            if resume.education:
                self._add_pdf_section(story, styles, "EDUCATION")
                for edu in resume.education:
                    self._add_pdf_education(story, styles, edu)

            # Skills section
            if resume.skills:
                self._add_pdf_section(story, styles, "TECHNICAL SKILLS")
                skills_text = " • ".join(resume.skills)
                skills_para = Paragraph(self._sanitize_text(skills_text), styles["Body"])
                story.append(skills_para)
                story.append(Spacer(1, 0.2 * inch))

            # Projects section (if present)
            if resume.projects:
                self._add_pdf_section(story, styles, "PROJECTS")
                for project in resume.projects:
                    self._add_pdf_project(story, styles, project)

            # Certifications section (if present)
            if resume.certifications:
                self._add_pdf_section(story, styles, "CERTIFICATIONS")
                for cert in resume.certifications:
                    self._add_pdf_certification(story, styles, cert)

            # Build PDF
            doc.build(story)

            # Get PDF bytes
            pdf_bytes = buffer.getvalue()
            buffer.close()

            logger.info(f"Successfully generated PDF ({len(pdf_bytes)} bytes)")
            return pdf_bytes

        except Exception as e:
            logger.error(f"Failed to generate PDF: {str(e)}")
            raise

    def _create_pdf_styles(self) -> dict:
        """Create custom PDF styles for ATS-friendly formatting"""
        styles = getSampleStyleSheet()

        # Name style
        if "Name" not in styles:
            styles.add(
                ParagraphStyle(
                    name="Name",
                    parent=styles["Heading1"],
                    fontSize=self.FONT_SIZE_NAME,
                    textColor=self.COLOR_PRIMARY,
                    spaceAfter=6,
                    alignment=1,  # Center
                    fontName=f"{self.FONT_NAME}-Bold",
                )
            )

        # Contact info style
        if "Contact" not in styles:
            styles.add(
                ParagraphStyle(
                    name="Contact",
                    parent=styles["Normal"],
                    fontSize=self.FONT_SIZE_SMALL,
                    textColor=self.COLOR_PRIMARY,
                    spaceAfter=12,
                    alignment=1,  # Center
                    fontName=self.FONT_NAME,
                )
            )

        # Section header style
        if "SectionHeader" not in styles:
            styles.add(
                ParagraphStyle(
                    name="SectionHeader",
                    parent=styles["Heading2"],
                    fontSize=self.FONT_SIZE_SECTION,
                    textColor=self.COLOR_SECTION,
                    spaceAfter=8,
                    spaceBefore=12,
                    fontName=f"{self.FONT_NAME}-Bold",
                    borderWidth=1,
                    borderColor=self.COLOR_SECTION,
                    borderPadding=4,
                    borderRadius=0,
                )
            )

        # Job title style
        if "JobTitle" not in styles:
            styles.add(
                ParagraphStyle(
                    name="JobTitle",
                    parent=styles["Normal"],
                    fontSize=self.FONT_SIZE_BODY,
                    textColor=self.COLOR_PRIMARY,
                    spaceAfter=2,
                    fontName=f"{self.FONT_NAME}-Bold",
                )
            )

        # Company style
        if "Company" not in styles:
            styles.add(
                ParagraphStyle(
                    name="Company",
                    parent=styles["Normal"],
                    fontSize=self.FONT_SIZE_BODY,
                    textColor=self.COLOR_PRIMARY,
                    spaceAfter=4,
                    fontName=f"{self.FONT_NAME}-Oblique",
                )
            )

        # Body text style
        if "Body" not in styles:
            styles.add(
                ParagraphStyle(
                    name="Body",
                    parent=styles["Normal"],
                    fontSize=self.FONT_SIZE_BODY,
                    textColor=self.COLOR_PRIMARY,
                    leading=self.FONT_SIZE_BODY * self.LINE_SPACING,
                    spaceAfter=4,
                    fontName=self.FONT_NAME,
                )
            )

        # Bullet style - use custom name to avoid conflict
        if "ResumeBullet" not in styles:
            styles.add(
                ParagraphStyle(
                    name="ResumeBullet",
                    parent=styles["Normal"],
                    fontSize=self.FONT_SIZE_BODY,
                    textColor=self.COLOR_PRIMARY,
                    leading=self.FONT_SIZE_BODY * self.LINE_SPACING,
                    spaceAfter=3,
                    leftIndent=20,
                    bulletIndent=10,
                    fontName=self.FONT_NAME,
                )
            )

        return styles

    def _add_pdf_header(self, story: List, styles: dict, resume: Resume):
        """Add header section with name and contact info"""
        info = resume.personalInfo

        # Name
        name_para = Paragraph(self._sanitize_text(info.name), styles["Name"])
        story.append(name_para)

        # Contact information
        contact_parts = []
        if info.email:
            contact_parts.append(info.email)
        if info.phone:
            contact_parts.append(info.phone)
        if info.location:
            contact_parts.append(info.location)

        if contact_parts:
            contact_text = " • ".join(contact_parts)
            contact_para = Paragraph(self._sanitize_text(contact_text), styles["Contact"])
            story.append(contact_para)

        # Links
        link_parts = []
        if info.linkedin:
            link_parts.append(self._clean_url(info.linkedin))
        if info.github:
            link_parts.append(self._clean_url(info.github))
        if info.website:
            link_parts.append(self._clean_url(info.website))

        if link_parts:
            link_text = " • ".join(link_parts)
            link_para = Paragraph(
                f'<font color="{self.COLOR_LINK}">{self._sanitize_text(link_text)}</font>',
                styles["Contact"],
            )
            story.append(link_para)

        story.append(Spacer(1, 0.15 * inch))

    def _add_pdf_section(self, story: List, styles: dict, title: str):
        """Add a section header"""
        section_para = Paragraph(title, styles["SectionHeader"])
        story.append(section_para)

    def _add_pdf_experience(self, story: List, styles: dict, exp):
        """Add work experience entry"""
        # Position and Company
        position_para = Paragraph(self._sanitize_text(exp.position), styles["JobTitle"])
        story.append(position_para)

        company_location = exp.company
        if exp.location:
            company_location += f" - {exp.location}"

        company_para = Paragraph(self._sanitize_text(company_location), styles["Company"])
        story.append(company_para)

        # Date range
        date_range = self._format_date_range(exp.startDate, exp.endDate)
        date_para = Paragraph(
            f'<font size="{self.FONT_SIZE_SMALL}">{date_range}</font>', styles["Body"]
        )
        story.append(date_para)
        story.append(Spacer(1, 0.05 * inch))

        # Responsibilities/achievements
        for desc in exp.description:
            bullet_text = f"• {self._sanitize_text(desc)}"
            bullet_para = Paragraph(bullet_text, styles["ResumeBullet"])
            story.append(bullet_para)

        story.append(Spacer(1, 0.15 * inch))

    def _add_pdf_education(self, story: List, styles: dict, edu):
        """Add education entry"""
        # Degree and field
        degree_text = edu.degree
        if edu.field:
            degree_text += f" in {edu.field}"

        degree_para = Paragraph(self._sanitize_text(degree_text), styles["JobTitle"])
        story.append(degree_para)

        # Institution
        institution_para = Paragraph(self._sanitize_text(edu.institution), styles["Company"])
        story.append(institution_para)

        # Date and GPA
        date_range = self._format_date_range(edu.startDate, edu.endDate)
        date_text = date_range
        if edu.gpa:
            date_text += f" • GPA: {edu.gpa}"

        date_para = Paragraph(
            f'<font size="{self.FONT_SIZE_SMALL}">{self._sanitize_text(date_text)}</font>',
            styles["Body"],
        )
        story.append(date_para)

        # Achievements
        if edu.achievements:
            story.append(Spacer(1, 0.05 * inch))
            for achievement in edu.achievements:
                bullet_text = f"• {self._sanitize_text(achievement)}"
                bullet_para = Paragraph(bullet_text, styles["ResumeBullet"])
                story.append(bullet_para)

        story.append(Spacer(1, 0.15 * inch))

    def _add_pdf_project(self, story: List, styles: dict, project):
        """Add project entry"""
        # Project name
        project_name = project.name
        if project.link:
            project_name += f" ({self._clean_url(project.link)})"

        project_para = Paragraph(self._sanitize_text(project_name), styles["JobTitle"])
        story.append(project_para)

        # Technologies
        if project.technologies:
            tech_text = f"Technologies: {', '.join(project.technologies)}"
            tech_para = Paragraph(
                f'<font size="{self.FONT_SIZE_SMALL}">{self._sanitize_text(tech_text)}</font>',
                styles["Body"],
            )
            story.append(tech_para)

        # Description
        if project.description:
            story.append(Spacer(1, 0.05 * inch))
            desc_para = Paragraph(self._sanitize_text(project.description), styles["Body"])
            story.append(desc_para)

        # Highlights
        if project.highlights:
            story.append(Spacer(1, 0.05 * inch))
            for highlight in project.highlights:
                bullet_text = f"• {self._sanitize_text(highlight)}"
                bullet_para = Paragraph(bullet_text, styles["ResumeBullet"])
                story.append(bullet_para)

        story.append(Spacer(1, 0.15 * inch))

    def _add_pdf_certification(self, story: List, styles: dict, cert):
        """Add certification entry"""
        # Certification name
        cert_para = Paragraph(self._sanitize_text(cert.name), styles["JobTitle"])
        story.append(cert_para)

        # Issuer and date
        issuer_text = f"{cert.issuer} • {cert.date}"
        if cert.credentialId:
            issuer_text += f" • ID: {cert.credentialId}"

        issuer_para = Paragraph(
            f'<font size="{self.FONT_SIZE_SMALL}">{self._sanitize_text(issuer_text)}</font>',
            styles["Body"],
        )
        story.append(issuer_para)

        story.append(Spacer(1, 0.1 * inch))

    # ==================== DOCX GENERATION ====================

    def generate_docx(self, resume: Resume) -> bytes:
        """
        Generate an ATS-friendly DOCX resume.

        Args:
            resume: Resume model from Gemini tailoring

        Returns:
            DOCX file as bytes

        Raises:
            Exception: If DOCX generation fails
        """
        try:
            logger.info(
                f"Generating DOCX for resume: {resume.name} (candidate: {resume.personalInfo.name})"
            )

            # Create document
            doc = Document()

            # Set document properties
            doc.core_properties.title = f"{resume.personalInfo.name} - Resume"
            doc.core_properties.author = resume.personalInfo.name

            # Set margins (0.75 inches)
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.75)
                section.bottom_margin = Inches(0.75)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)

            # Header section
            self._add_docx_header(doc, resume)

            # Summary section (if present)
            if resume.personalInfo.summary:
                self._add_docx_section(doc, "PROFESSIONAL SUMMARY")
                para = doc.add_paragraph(resume.personalInfo.summary)
                self._format_docx_body(para)
                doc.add_paragraph()  # Spacing

            # Experience section
            if resume.experience:
                self._add_docx_section(doc, "PROFESSIONAL EXPERIENCE")
                for exp in resume.experience:
                    self._add_docx_experience(doc, exp)

            # Education section
            if resume.education:
                self._add_docx_section(doc, "EDUCATION")
                for edu in resume.education:
                    self._add_docx_education(doc, edu)

            # Skills section
            if resume.skills:
                self._add_docx_section(doc, "TECHNICAL SKILLS")
                skills_text = " • ".join(resume.skills)
                para = doc.add_paragraph(skills_text)
                self._format_docx_body(para)
                doc.add_paragraph()  # Spacing

            # Projects section (if present)
            if resume.projects:
                self._add_docx_section(doc, "PROJECTS")
                for project in resume.projects:
                    self._add_docx_project(doc, project)

            # Certifications section (if present)
            if resume.certifications:
                self._add_docx_section(doc, "CERTIFICATIONS")
                for cert in resume.certifications:
                    self._add_docx_certification(doc, cert)

            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            docx_bytes = buffer.getvalue()
            buffer.close()

            logger.info(f"Successfully generated DOCX ({len(docx_bytes)} bytes)")
            return docx_bytes

        except Exception as e:
            logger.error(f"Failed to generate DOCX: {str(e)}")
            raise

    def _add_docx_header(self, doc: Document, resume: Resume):
        """Add header section with name and contact info"""
        info = resume.personalInfo

        # Name
        name_para = doc.add_paragraph(info.name)
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.runs[0]
        name_run.font.size = Pt(18)
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(26, 26, 26)

        # Contact information
        contact_parts = []
        if info.email:
            contact_parts.append(info.email)
        if info.phone:
            contact_parts.append(info.phone)
        if info.location:
            contact_parts.append(info.location)

        if contact_parts:
            contact_text = " • ".join(contact_parts)
            contact_para = doc.add_paragraph(contact_text)
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_run = contact_para.runs[0]
            contact_run.font.size = Pt(10)
            contact_run.font.color.rgb = RGBColor(26, 26, 26)

        # Links
        link_parts = []
        if info.linkedin:
            link_parts.append(self._clean_url(info.linkedin))
        if info.github:
            link_parts.append(self._clean_url(info.github))
        if info.website:
            link_parts.append(self._clean_url(info.website))

        if link_parts:
            link_text = " • ".join(link_parts)
            link_para = doc.add_paragraph(link_text)
            link_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            link_run = link_para.runs[0]
            link_run.font.size = Pt(10)
            link_run.font.color.rgb = RGBColor(29, 78, 216)  # Blue color for links

        doc.add_paragraph()  # Spacing

    def _add_docx_section(self, doc: Document, title: str):
        """Add a section header"""
        para = doc.add_paragraph(title)
        run = para.runs[0]
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(37, 99, 235)  # Professional blue

    def _add_docx_experience(self, doc: Document, exp):
        """Add work experience entry"""
        # Position
        position_para = doc.add_paragraph(exp.position)
        position_run = position_para.runs[0]
        position_run.font.size = Pt(11)
        position_run.font.bold = True

        # Company and location
        company_text = exp.company
        if exp.location:
            company_text += f" - {exp.location}"

        company_para = doc.add_paragraph(company_text)
        company_run = company_para.runs[0]
        company_run.font.size = Pt(11)
        company_run.font.italic = True

        # Date range
        date_range = self._format_date_range(exp.startDate, exp.endDate)
        date_para = doc.add_paragraph(date_range)
        date_run = date_para.runs[0]
        date_run.font.size = Pt(10)

        # Responsibilities/achievements
        for desc in exp.description:
            para = doc.add_paragraph(desc, style="List Bullet")
            self._format_docx_body(para)

        doc.add_paragraph()  # Spacing

    def _add_docx_education(self, doc: Document, edu):
        """Add education entry"""
        # Degree and field
        degree_text = edu.degree
        if edu.field:
            degree_text += f" in {edu.field}"

        degree_para = doc.add_paragraph(degree_text)
        degree_run = degree_para.runs[0]
        degree_run.font.size = Pt(11)
        degree_run.font.bold = True

        # Institution
        institution_para = doc.add_paragraph(edu.institution)
        institution_run = institution_para.runs[0]
        institution_run.font.size = Pt(11)
        institution_run.font.italic = True

        # Date and GPA
        date_range = self._format_date_range(edu.startDate, edu.endDate)
        date_text = date_range
        if edu.gpa:
            date_text += f" • GPA: {edu.gpa}"

        date_para = doc.add_paragraph(date_text)
        date_run = date_para.runs[0]
        date_run.font.size = Pt(10)

        # Achievements
        if edu.achievements:
            for achievement in edu.achievements:
                para = doc.add_paragraph(achievement, style="List Bullet")
                self._format_docx_body(para)

        doc.add_paragraph()  # Spacing

    def _add_docx_project(self, doc: Document, project):
        """Add project entry"""
        # Project name
        project_name = project.name
        if project.link:
            project_name += f" ({self._clean_url(project.link)})"

        project_para = doc.add_paragraph(project_name)
        project_run = project_para.runs[0]
        project_run.font.size = Pt(11)
        project_run.font.bold = True

        # Technologies
        if project.technologies:
            tech_text = f"Technologies: {', '.join(project.technologies)}"
            tech_para = doc.add_paragraph(tech_text)
            tech_run = tech_para.runs[0]
            tech_run.font.size = Pt(10)

        # Description
        if project.description:
            desc_para = doc.add_paragraph(project.description)
            self._format_docx_body(desc_para)

        # Highlights
        if project.highlights:
            for highlight in project.highlights:
                para = doc.add_paragraph(highlight, style="List Bullet")
                self._format_docx_body(para)

        doc.add_paragraph()  # Spacing

    def _add_docx_certification(self, doc: Document, cert):
        """Add certification entry"""
        # Certification name
        cert_para = doc.add_paragraph(cert.name)
        cert_run = cert_para.runs[0]
        cert_run.font.size = Pt(11)
        cert_run.font.bold = True

        # Issuer and date
        issuer_text = f"{cert.issuer} • {cert.date}"
        if cert.credentialId:
            issuer_text += f" • ID: {cert.credentialId}"

        issuer_para = doc.add_paragraph(issuer_text)
        issuer_run = issuer_para.runs[0]
        issuer_run.font.size = Pt(10)

        doc.add_paragraph()  # Spacing

    def _format_docx_body(self, para):
        """Apply standard body formatting to a paragraph"""
        for run in para.runs:
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(26, 26, 26)

    # ==================== UTILITY METHODS ====================

    def _format_date_range(self, start_date: str, end_date: str) -> str:
        """
        Format date range for display.

        Args:
            start_date: Start date (YYYY-MM or YYYY)
            end_date: End date (YYYY-MM, YYYY, or "Present")

        Returns:
            Formatted date range string
        """
        if not start_date:
            return ""

        # Handle "Present" end date
        if end_date and end_date.lower() == "present":
            return f"{start_date} - Present"

        # Handle regular date range
        if end_date:
            return f"{start_date} - {end_date}"

        return start_date

    def _clean_url(self, url: str) -> str:
        """
        Clean URL for display (remove https://, www., etc.)

        Args:
            url: Full URL

        Returns:
            Cleaned URL for display
        """
        if not url:
            return ""

        # Remove protocol
        cleaned = url.replace("https://", "").replace("http://", "")

        # Remove www.
        cleaned = cleaned.replace("www.", "")

        # Remove trailing slash
        if cleaned.endswith("/"):
            cleaned = cleaned[:-1]

        return cleaned

    def _sanitize_text(self, text: str) -> str:
        """
        Sanitize text for PDF generation (escape special characters).

        Args:
            text: Raw text

        Returns:
            Sanitized text safe for PDF rendering
        """
        if not text:
            return ""

        # Escape XML special characters for ReportLab
        text = text.replace("&", "&amp;")
        text = text.replace("<", "&lt;")
        text = text.replace(">", "&gt;")

        # Remove problematic characters
        text = text.replace("\n", " ")
        text = text.replace("\r", "")
        text = text.replace("\t", " ")

        # Collapse multiple spaces
        while "  " in text:
            text = text.replace("  ", " ")

        return text.strip()


# Singleton instance
_document_generator: Optional[DocumentGeneratorService] = None


def get_document_generator() -> DocumentGeneratorService:
    """
    Get singleton instance of DocumentGeneratorService.

    Returns:
        DocumentGeneratorService instance
    """
    global _document_generator
    if _document_generator is None:
        _document_generator = DocumentGeneratorService()
    return _document_generator
