"""
Template-based document generator for ATS-friendly PDF and DOCX resumes.
Uses HTML templates with Jinja2 for consistent, professional formatting.

This replaces the manual ReportLab-based generation with a template-driven approach:
- HTML templates → WeasyPrint → PDF (ATS-optimized)
- Template data → python-docx → DOCX (Native Generation with Tables)
- High visual consistency between PDF and DOCX
"""

from __future__ import annotations

import io
import os
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from jinja2 import Environment, FileSystemLoader, select_autoescape
from loguru import logger

# Lazy-load WeasyPrint with Windows GTK fallback
try:
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
except OSError as e:
    WEASYPRINT_AVAILABLE = False
    WEASYPRINT_ERROR = str(e)
    # Create dummy classes to prevent import errors
    class HTML:
        def __init__(self, *args, **kwargs):
            raise RuntimeError(
                f"WeasyPrint is not properly installed. Missing system dependencies: {WEASYPRINT_ERROR}\n"
                "On Windows, install GTK+ runtime: https://github.com/tschoonj/GTK-for-Windows-Runtime-Environment-Installer/releases\n"
                "Or use Docker: 'make dev'"
            )
    class CSS:
        def __init__(self, *args, **kwargs):
            pass

try:
    from ..models.resume import Resume
except ImportError:
    from models.resume import Resume


class TemplateDocumentGenerator:
    """
    Template-based document generator for PDF and DOCX resumes.
    """

    def __init__(self, template_dir: Optional[str] = None):
        if template_dir is None:
            # Default to backend/templates directory
            current_file = Path(__file__).resolve()
            backend_root = current_file.parent.parent.parent  # Go up to backend/
            template_dir = str(backend_root / "templates")

        self.template_dir = template_dir

        self.jinja_env = Environment(
            loader=FileSystemLoader(template_dir),
            autoescape=select_autoescape(['html', 'xml']),
            trim_blocks=True,
            lstrip_blocks=True
        )

        logger.info(f"Initialized TemplateDocumentGenerator with templates from: {template_dir}")

    def _categorize_skills(self, skills: List[str]) -> Dict[str, List[str]]:
        categories = {
            "Languages": ["python", "javascript", "typescript", "java", "c++", "c#", "go", "rust", "ruby", "php", "kotlin", "swift", "scala"],
            "Frontend": ["react", "vue", "angular", "svelte", "html", "css", "sass", "tailwind", "bootstrap", "webpack", "vite"],
            "Backend": ["node.js", "express", "fastapi", "django", "flask", "spring", "asp.net", ".net", "rails"],
            "Databases": ["postgresql", "mysql", "mongodb", "redis", "elasticsearch", "sql", "nosql", "dynamodb", "cassandra"],
            "Cloud & DevOps": ["aws", "azure", "gcp", "docker", "kubernetes", "terraform", "jenkins", "github actions", "ci/cd", "linux"],
            "Tools & Other": []
        }

        categorized = {cat: [] for cat in categories.keys()}
        uncategorized = []

        for skill in skills:
            skill_lower = skill.lower()
            matched = False
            for category, keywords in categories.items():
                if category == "Tools & Other": continue
                for keyword in keywords:
                    if keyword in skill_lower or skill_lower in keyword:
                        categorized[category].append(skill)
                        matched = True
                        break
                if matched: break
            if not matched: uncategorized.append(skill)

        categorized["Tools & Other"] = uncategorized
        categorized = {k: v for k, v in categorized.items() if v}
        return categorized

    def _prepare_template_data(self, resume: Resume) -> Dict[str, Any]:
        data = resume.model_dump()

        if resume.skills and len(resume.skills) > 0:
            data['skillsGrouped'] = self._categorize_skills(resume.skills)

        for exp in data.get('experience', []):
            if exp.get('endDate') and exp['endDate'].lower() in ['present', 'current']:
                exp['endDate'] = 'Present'
            if 'description' in exp:
                exp['bullets'] = exp['description']

        for edu in data.get('education', []):
            if 'achievements' in edu and edu['achievements']:
                edu['coursework'] = ", ".join(edu['achievements'])

        data['generatedDate'] = 'Generated with Resume Tailor AI'

        # Env settings fallback
        try:
            from ..app.config import settings
            env_accent_color = settings.RESUME_ACCENT_COLOR
            env_font_family = settings.RESUME_FONT_FAMILY
            env_font_size = settings.RESUME_FONT_SIZE
        except (ImportError, ModuleNotFoundError):
            env_accent_color = '#1e3a5f'
            env_font_family = 'Roboto'
            env_font_size = 9
        
        # Use resume-specific accent color if provided, otherwise use env/default
        if resume.accentColor:
            logger.info(f"Using resume-specific accent color: {resume.accentColor}")
            data['accentColor'] = resume.accentColor
        else:
            data['accentColor'] = env_accent_color
        
        data['fontFamily'] = env_font_family
        data['fontSize'] = env_font_size

        return data

    def generate_pdf(self, resume: Resume, template_name: str = 'resume_template_professional.html') -> bytes:
        if not WEASYPRINT_AVAILABLE:
             raise RuntimeError(f"WeasyPrint unavailable: {WEASYPRINT_ERROR}")
        
        try:
            logger.info(f"Generating PDF for: {resume.personalInfo.name}")
            template_data = self._prepare_template_data(resume)
            template = self.jinja_env.get_template(template_name)
            html_content = template.render(**template_data)
            html = HTML(string=html_content, base_url=self.template_dir)
            pdf_bytes = html.write_pdf()
            logger.info(f"PDF generated successfully: {len(pdf_bytes)} bytes")
            return pdf_bytes
        except Exception as e:
            logger.error(f"Failed to generate PDF via WeasyPrint: {str(e)}")
            # Fallback logic here if needed, or re-raise
            raise e

    # ==================== DOCX GENERATION (Native) ====================

    def generate_docx(self, resume: Resume) -> bytes:
        """
        Generate ATS-friendly DOCX resume using python-docx.
        Uses native Word tables and styles to mirror the PDF layout.
        """
        try:
            logger.info(f"Generating DOCX for: {resume.personalInfo.name}")

            # Prepare data
            data = self._prepare_template_data(resume)
            accent_color_hex = data.get('accentColor', '#1e3a5f')
            accent_rgb = self._hex_to_rgb(accent_color_hex)
            
            # Create document
            doc = Document()
            
            # Set margins (Narrow: 0.5")
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.7)
                section.right_margin = Inches(0.7)

            # --- Header ---
            self._add_docx_header(doc, resume.personalInfo, accent_rgb)

            # --- Summary ---
            if resume.personalInfo.summary:
                self._add_docx_section_title(doc, "PROFESSIONAL SUMMARY", accent_rgb)
                para = doc.add_paragraph(resume.personalInfo.summary)
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                self._set_font(para, size=10)

            # --- Education ---
            if resume.education:
                self._add_docx_section_title(doc, "EDUCATION", accent_rgb)
                self._add_docx_education(doc, resume.education)

            # --- Skills ---
            if resume.skills:
                self._add_docx_section_title(doc, "SKILLS", accent_rgb)
                self._add_docx_skills(doc, resume.skills)

            # --- Experience ---
            if resume.experience:
                self._add_docx_section_title(doc, "WORK EXPERIENCE", accent_rgb)
                self._add_docx_experience(doc, resume.experience)

            # --- Projects ---
            if resume.projects:
                self._add_docx_section_title(doc, "PROJECTS", accent_rgb)
                self._add_docx_projects(doc, resume.projects, accent_rgb)

            # --- Certifications ---
            if resume.certifications:
                self._add_docx_section_title(doc, "CERTIFICATIONS", accent_rgb)
                self._add_docx_certifications(doc, resume.certifications)

            # Save to bytes
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            docx_bytes = buffer.read()

            logger.info(f"DOCX generated successfully: {len(docx_bytes)} bytes")
            return docx_bytes

        except Exception as e:
            logger.error(f"Failed to generate DOCX: {str(e)}")
            raise Exception(f"DOCX generation failed: {str(e)}")

    # ==================== DOCX HELPER METHODS ====================

    def _hex_to_rgb(self, hex_color: str) -> RGBColor:
        hex_color = hex_color.lstrip('#')
        try:
            return RGBColor(
                int(hex_color[0:2], 16),
                int(hex_color[2:4], 16),
                int(hex_color[4:6], 16)
            )
        except ValueError:
            return RGBColor(30, 58, 95)  # Default Fallback

    def _set_font(self, para, name='Arial', size=10, bold=False, italic=False, color=None):
        for run in para.runs:
            run.font.name = name
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.italic = italic
            if color: run.font.color.rgb = color
            
    def _add_docx_header(self, doc: Document, personal_info, accent_rgb):
        # Name
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(personal_info.name.upper())
        name_run.font.name = 'Arial'
        name_run.font.size = Pt(24)
        name_run.font.bold = True
        name_run.font.color.rgb = accent_rgb
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_para.paragraph_format.space_after = Pt(4)

        # Contact info line
        contact_parts = []
        if personal_info.email: contact_parts.append(personal_info.email)
        if personal_info.phone: contact_parts.append(personal_info.phone)
        if personal_info.location: contact_parts.append(personal_info.location)
        if personal_info.linkedin: 
            linkedin = personal_info.linkedin.replace('https://', '').replace('http://', '').replace('www.', '')
            contact_parts.append(linkedin)
        if personal_info.website:
            website = personal_info.website.replace('https://', '').replace('http://', '').replace('www.', '')
            contact_parts.append(website)
        
        if contact_parts:
            contact_para = doc.add_paragraph(" | ".join(contact_parts))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in contact_para.runs:
                run.font.size = Pt(9)
                run.font.name = 'Arial'

        doc.add_paragraph()  # Spacing

    def _add_docx_section_title(self, doc: Document, title: str, accent_rgb):
        para = doc.add_paragraph()
        run = para.add_run(title)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = accent_rgb
        
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(3)
        
        # Add a bottom border
        p = para._p
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _create_two_col_row(self, doc, left_text, right_text, bold=False):
        """Simulate a left-right layout using a 2-column table"""
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.allow_autofit = False
        
        # Widths: 5.0" left, 2.0" right roughly
        table.columns[0].width = Inches(5.0)
        table.columns[1].width = Inches(2.0)
        
        cell_left = table.cell(0, 0)
        cell_right = table.cell(0, 1)
        
        # Left Content
        p_left = cell_left.paragraphs[0]
        run_left = p_left.add_run(left_text)
        run_left.font.name = 'Arial'
        run_left.font.size = Pt(11)
        if bold: run_left.font.bold = True
        
        # Right Content
        p_right = cell_right.paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if right_text:
            run_right = p_right.add_run(right_text)
            run_right.font.name = 'Arial'
            run_right.font.size = Pt(10)
            run_right.font.italic = True
        
        # Format table to look invisible
        # This usually defaults to no borders in python-docx unless style is set

    def _add_docx_experience(self, doc: Document, experience: List):
        for exp in experience:
            # Header Row: Company & Dates
            date_str = f"{exp.startDate}"
            if exp.endDate:
                date_str += f" - {exp.endDate}"
            
            self._create_two_col_row(doc, exp.company, date_str, bold=True)
            
            # Sub-header Row: Position & Location
            self._create_two_col_row(doc, exp.position, exp.location or "")
            
            # Bullets
            if exp.description:
                for bullet in exp.description:
                    p = doc.add_paragraph(style='List Bullet')
                    p.paragraph_format.left_indent = Inches(0.2)
                    p.paragraph_format.space_after = Pt(0)
                    run = p.add_run(bullet)
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)

            doc.add_paragraph() # Spacing

    def _add_docx_education(self, doc: Document, education: List):
        for edu in education:
            date_str = f"{edu.startDate} - {edu.endDate}" if edu.endDate else edu.startDate
            header_text = f"{edu.institution}"
            # if edu.location: header_text += f", {edu.location}"
            
            self._create_two_col_row(doc, header_text, date_str, bold=True)
            
            # Degree line
            degree_str = edu.degree
            if edu.field: degree_str += f" in {edu.field}"
            if edu.gpa: degree_str += f" (GPA: {edu.gpa})"
            
            p = doc.add_paragraph(degree_str)
            p.paragraph_format.space_after = Pt(4)
            p.runs[0].font.name = 'Arial'
            p.runs[0].font.size = Pt(10)
            
            if edu.achievements:
                p_cw = doc.add_paragraph(f"Relevant Coursework: {', '.join(edu.achievements)}")
                p_cw.paragraph_format.space_after = Pt(8)
                self._set_font(p_cw, size=10)

    def _add_docx_projects(self, doc: Document, projects: List, accent_rgb):
        for proj in projects:
            date_str = f"{proj.startDate} - {proj.endDate}" if proj.endDate else proj.startDate
            self._create_two_col_row(doc, proj.name, date_str, bold=True)
            
            # Link
            if proj.link:
                p = doc.add_paragraph()
                r = p.add_run(proj.link)
                r.font.color.rgb = accent_rgb
                r.font.underline = True
                r.font.size = Pt(9)
                p.paragraph_format.space_after = Pt(2)

            # Tech
            if proj.technologies:
                p = doc.add_paragraph(f"Stack: {', '.join(proj.technologies)}")
                p.paragraph_format.space_after = Pt(4)
                if p.runs:
                    p.runs[0].font.italic = True
                    p.runs[0].font.size = Pt(10)
                
            # Bullets
            if proj.highlights:
                for highlight in proj.highlights:
                    p = doc.add_paragraph(style='List Bullet')
                    p.paragraph_format.left_indent = Inches(0.2)
                    p.paragraph_format.space_after = Pt(0)
                    run = p.add_run(highlight)
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)

            doc.add_paragraph()

    def _add_docx_skills(self, doc: Document, skills: List[str]):
        categorized = self._categorize_skills(skills)
        if categorized:
            for category, skill_list in categorized.items():
                p = doc.add_paragraph()
                r_cat = p.add_run(f"{category}: ")
                r_cat.font.bold = True
                r_cat.font.name = 'Arial'
                r_cat.font.size = Pt(10)
                
                r_list = p.add_run(", ".join(skill_list))
                r_list.font.name = 'Arial'
                r_list.font.size = Pt(10)
        else:
             p = doc.add_paragraph(", ".join(skills))
             self._set_font(p, size=10)
        
    def _add_docx_certifications(self, doc: Document, certifications: List):
        for cert in certifications:
            date_str = f"({cert.date})" if cert.date else ""
            self._create_two_col_row(doc, f"{cert.name} - {cert.issuer}", date_str)

def get_template_generator(template_dir: Optional[str] = None) -> TemplateDocumentGenerator:
    return TemplateDocumentGenerator(template_dir=template_dir)
